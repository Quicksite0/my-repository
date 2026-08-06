import tkinter as tk
from tkinter import ttk, font, filedialog, messagebox
import os
import time
import json
import io
import base64

# check for python-docx
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# check for pillow
try:
    from PIL import Image, ImageTk, ImageGrab
    PIL_AVAILABLE = True
    try:
        LANCZOS = Image.Resampling.LANCZOS
    except AttributeError:
        try:
            LANCZOS = Image.ANTIALIAS
        except AttributeError:
            LANCZOS = Image.BICUBIC
except ImportError:
    PIL_AVAILABLE = False

class ImageWidget(tk.Frame):
    def __init__(self, parent, editor, image, app, width=None, height=None):
        # determine theme background
        bg_color = "#1e1e1e" if app.is_dark_mode else "white"
        super().__init__(parent, bd=1, relief=tk.FLAT, bg=bg_color)
        self.editor = editor
        self.original_image = image
        self.app = app
        self.width = width if width else image.width
        self.height = height if height else image.height
        
        self.aspect_ratio = image.width / image.height
        
        # limit starting image size
        if not width and not height:
            max_size = 300
            if self.width > max_size or self.height > max_size:
                if self.width > self.height:
                    self.width = max_size
                    self.height = int(max_size / self.aspect_ratio)
                else:
                    self.height = max_size
                    self.width = int(max_size * self.aspect_ratio)
        
        self.update_image()
        
        # resize handle in the bottom right corner
        self.handle = tk.Frame(self, width=8, height=8, bg="#555555", cursor="size_nw_se")
        self.handle.place(relx=1.0, rely=1.0, anchor="se")
        
        # event listeners
        self.label.bind("<Button-1>", self.on_select)
        self.label.bind("<B1-Motion>", self.on_drag)
        self.label.bind("<ButtonRelease-1>", self.on_release)
        
        self.handle.bind("<Button-1>", self.start_resize)
        self.handle.bind("<B1-Motion>", self.do_resize)
        
        # handle delete focus
        self.label.config(takefocus=True)
        self.label.bind("<Delete>", self.on_delete_key)
        self.label.bind("<BackSpace>", self.on_delete_key)
        self.label.bind("<FocusOut>", self.on_deselect)
        
    def update_image(self):
        """resize and update image display"""
        resized = self.original_image.resize((self.width, self.height), LANCZOS)
        self.img_tk = ImageTk.PhotoImage(resized)
        if hasattr(self, 'label'):
            self.label.config(image=self.img_tk)
        else:
            self.label = tk.Label(self, image=self.img_tk, bg=self.cget("bg"))
            self.label.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.config(width=self.width, height=self.height)
        
    def on_select(self, event):
        """highlight image and set focus"""
        self.app.save_state()
        self.config(bd=2, bg="#007BFF")  # visual blue selection border
        self.label.focus_set()
        
    def on_deselect(self, event):
        """remove highlight border"""
        bg_color = "#1e1e1e" if self.app.is_dark_mode else "white"
        self.config(bd=1, bg=bg_color, relief=tk.FLAT)
        
    def on_delete_key(self, event):
        """delete image on keystroke"""
        self.app.save_state()
        try:
            idx = self.editor.index(self)
            self.editor.delete(idx)
        except Exception:
            pass
        self.app.update_live_stats()
        
    def start_resize(self, event):
        """save initial resize sizes"""
        self.app.save_state()
        self.resize_start_x = event.x_root
        self.resize_start_y = event.y_root
        self.resize_start_w = self.width
        self.resize_start_h = self.height
        
    def do_resize(self, event):
        """adjust size with cursor movement"""
        dx = event.x_root - self.resize_start_x
        dy = event.y_root - self.resize_start_y
        
        # scale keeping aspect ratio
        if abs(dx) > abs(dy):
            new_w = max(30, self.resize_start_w + dx)
            new_h = int(new_w / self.aspect_ratio)
        else:
            new_h = max(30, self.resize_start_h + dy)
            new_w = int(new_h * self.aspect_ratio)
            
        self.width = new_w
        self.height = new_h
        self.update_image()
        
    def on_drag(self, event):
        """preview new cursor position on drag"""
        editor_x = event.x_root - self.editor.winfo_rootx()
        editor_y = event.y_root - self.editor.winfo_rooty()
        target_index = self.editor.index(f"@{editor_x},{editor_y}")
        self.editor.mark_set(tk.INSERT, target_index)
        
    def on_release(self, event):
        """move image to target cursor index"""
        self.on_deselect(None)
        editor_x = event.x_root - self.editor.winfo_rootx()
        editor_y = event.y_root - self.editor.winfo_rooty()
        target_index = self.editor.index(f"@{editor_x},{editor_y}")
        
        try:
            orig_index = self.editor.index(self)
            if self.editor.compare(target_index, "!=", orig_index):
                self.app.save_state()
                self.editor.delete(orig_index)
                self.editor.window_create(target_index, window=self)
        except Exception:
            pass
        self.app.update_live_stats()

class PyDocApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PyDoc Architect - Professional Word Processor")
        
        # maximize window
        try:
            self.root.state('zoomed')
        except Exception:
            # fallback window size
            self.root.geometry("1100x800")
        
        # app state
        self.current_file = None
        self.is_dark_mode = False
        self.format_clipboard = None  # stores copied format tags
        
        # undo and redo stacks
        self.undo_stack = []
        self.redo_stack = []
        self.last_keystroke_time = time.time()
        
        # set up styles and fonts
        self.setup_styles()
        
        # start main interface
        self.build_main_interface()
        
        # set initial stats
        self.update_live_stats()

    def setup_styles(self):
        """set up styles with fallbacks"""
        self.style = ttk.Style()
        self.style.theme_use('clam')
        
        # default font
        self.default_font_family = "Arial"
        self.default_font_size = 12

    def build_main_interface(self):
        """build toolbar, editor, and status bar"""
        # menu bar
        self.main_menu = tk.Menu(self.root)
        self.root.config(menu=self.main_menu)
        
        file_menu = tk.Menu(self.main_menu, tearoff=0)
        self.main_menu.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="New", command=self.new_file)
        file_menu.add_command(label="Open", command=self.open_file)
        file_menu.add_command(label="Save", command=self.save_file)
        file_menu.add_command(label="Save As...", command=self.save_as_file)
        file_menu.add_separator()
        file_menu.add_command(label="Undo", command=self.undo_action, accelerator="Ctrl+Z")
        file_menu.add_command(label="Redo", command=self.redo_action, accelerator="Ctrl+Y")
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)

        # toolbar
        self.ribbon = tk.Frame(self.root, bg="#f3f3f3", bd=1, relief=tk.RAISED)
        self.ribbon.pack(side=tk.TOP, fill=tk.X)

        # font controls
        self.font_frame = tk.LabelFrame(self.ribbon, text="Font & Style", padx=5, pady=5, bg="#f3f3f3")
        self.font_frame.pack(side=tk.LEFT, padx=5, pady=5)

        self.font_family = tk.StringVar(value="Arial")
        self.font_choice = ttk.Combobox(self.font_frame, textvariable=self.font_family, width=15)
        
        # get system fonts safely
        try:
            families = list(font.families())
            cleaned_families = sorted(list(set(str(f) for f in families if f and not f.startswith("@"))))
            self.font_choice['values'] = cleaned_families
        except Exception:
            self.font_choice['values'] = ["Arial", "Courier New", "Georgia", "Times New Roman", "Verdana"]
            
        self.font_choice.pack(side=tk.LEFT, padx=2)
        self.font_choice.bind("<<ComboboxSelected>>", self.change_font)

        self.font_size = tk.IntVar(value=12)
        self.size_choice = ttk.Spinbox(self.font_frame, from_=8, to=72, textvariable=self.font_size, width=5, command=self.change_font)
        self.size_choice.pack(side=tk.LEFT, padx=2)
        self.size_choice.bind("<Return>", self.change_font)

        # format buttons
        self.btn_bold = tk.Button(self.font_frame, text="B", font=("Arial", 10, "bold"), width=3, command=self.toggle_bold)
        self.btn_bold.pack(side=tk.LEFT, padx=2)
        
        self.btn_italic = tk.Button(self.font_frame, text="I", font=("Arial", 10, "italic"), width=3, command=self.toggle_italic)
        self.btn_italic.pack(side=tk.LEFT, padx=2)

        # list controls
        self.bullet_frame = tk.LabelFrame(self.ribbon, text="Lists", padx=5, pady=5, bg="#f3f3f3")
        self.bullet_frame.pack(side=tk.LEFT, padx=5, pady=5)

        self.bullet_type = tk.StringVar(value="•")
        self.bullet_choice = ttk.Combobox(self.bullet_frame, textvariable=self.bullet_type, width=5, state="readonly")
        self.bullet_choice['values'] = ("•", "○", "▪", "➢", "★")
        self.bullet_choice.pack(side=tk.LEFT, padx=2)

        self.btn_bullet = tk.Button(self.bullet_frame, text="Add Bullet", command=self.insert_bullet)
        self.btn_bullet.pack(side=tk.LEFT, padx=2)

        # history controls
        self.history_frame = tk.LabelFrame(self.ribbon, text="History", padx=5, pady=5, bg="#f3f3f3")
        self.history_frame.pack(side=tk.LEFT, padx=5, pady=5)

        self.btn_undo = tk.Button(self.history_frame, text="⟲ Undo", command=self.undo_action)
        self.btn_undo.pack(side=tk.LEFT, padx=2)

        self.btn_redo = tk.Button(self.history_frame, text="⟳ Redo", command=self.redo_action)
        self.btn_redo.pack(side=tk.LEFT, padx=2)

        # format copy, paste and image actions
        self.paint_frame = tk.LabelFrame(self.ribbon, text="Clipboard", padx=5, pady=5, bg="#f3f3f3")
        self.paint_frame.pack(side=tk.LEFT, padx=5, pady=5)

        self.btn_copy_fmt = tk.Button(self.paint_frame, text="Copy Format", command=self.copy_format)
        self.btn_copy_fmt.pack(side=tk.LEFT, padx=2)

        self.btn_paste_fmt = tk.Button(self.paint_frame, text="Paste Format", command=self.paste_format)
        self.btn_paste_fmt.pack(side=tk.LEFT, padx=2)

        self.btn_paste_img = tk.Button(self.paint_frame, text="Paste Image", command=self.paste_image)
        self.btn_paste_img.pack(side=tk.LEFT, padx=2)

        # settings
        self.config_frame = tk.LabelFrame(self.ribbon, text="Settings", padx=5, pady=5, bg="#f3f3f3")
        self.config_frame.pack(side=tk.LEFT, padx=5, pady=5)
        
        self.btn_dark_mode = tk.Button(self.config_frame, text="Dark Mode", command=self.toggle_dark_mode)
        self.btn_dark_mode.pack(side=tk.LEFT, padx=2)

        # stats panel
        self.stats_frame = tk.LabelFrame(self.ribbon, text="Document Stats", padx=10, pady=5, bg="#f3f3f3")
        self.stats_frame.pack(side=tk.RIGHT, padx=10, pady=5)
        
        self.live_word_label = tk.Label(self.stats_frame, text="Words: 0", font=("Arial", 10, "bold"), bg="#f3f3f3", fg="#007BFF")
        self.live_word_label.pack(side=tk.TOP)
        
        self.live_char_label = tk.Label(self.stats_frame, text="Chars: 0", font=("Arial", 8), bg="#f3f3f3", fg="#666666")
        self.live_char_label.pack(side=tk.TOP)

        # editor area
        self.text_frame = tk.Frame(self.root)
        self.text_frame.pack(expand=True, fill=tk.BOTH)

        self.scroll_y = tk.Scrollbar(self.text_frame)
        self.scroll_y.pack(side=tk.RIGHT, fill=tk.Y)

        self.editor = tk.Text(self.text_frame, 
                              undo=False,  # disable native undo for custom implementation
                              font=(self.default_font_family, self.default_font_size),
                              yscrollcommand=self.scroll_y.set,
                              padx=20, pady=20,
                              wrap=tk.WORD)
        self.editor.pack(expand=True, fill=tk.BOTH)
        self.scroll_y.config(command=self.editor.yview)

        # key bindings
        self.editor.bind("<KeyRelease>", self.update_live_stats)
        self.editor.bind("<<Modified>>", self.update_live_stats)
        self.editor.bind("<KeyPress>", self.on_key_press)
        
        # override native undo keys
        self.editor.bind("<Control-z>", self.undo_action)
        self.editor.bind("<Control-y>", self.redo_action)
        self.editor.bind("<Control-Z>", self.undo_action)
        self.editor.bind("<Control-Y>", self.redo_action)
        
        # override custom paste handler
        self.editor.bind("<Control-v>", self.on_paste_event)
        self.editor.bind("<Control-V>", self.on_paste_event)

        # status bar
        self.status_bar = tk.Label(self.root, text="Ready", anchor=tk.W, relief=tk.SUNKEN)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

    # live stats
    def update_live_stats(self, event=None):
        """update word and char counts"""
        content = self.editor.get(1.0, "end-1c").strip()
        
        # count words
        words = len(content.split()) if content else 0
        chars = len(content)
        
        self.live_word_label.config(text=f"Words: {words}")
        self.live_char_label.config(text=f"Chars: {chars}")
        
        # reset modify flag
        self.editor.edit_modified(False)

    # style tags
    def get_or_create_style_tag(self, family, size, weight, slant):
        """get or make style tag"""
        tag_name = f"style_{family}_{size}_{weight}_{slant}"
        if tag_name not in self.editor.tag_names():
            f = font.Font(family=family, size=size, weight=weight, slant=slant)
            self.editor.tag_configure(tag_name, font=f)
        return tag_name

    def apply_formatting(self, font_family=None, font_size=None, bold=None, italic=None):
        """apply formatting to selection"""
        if not self.editor.tag_ranges(tk.SEL):
            # if no selection, set global defaults
            if font_family:
                self.default_font_family = font_family
            if font_size:
                self.default_font_size = font_size
            self.editor.configure(font=(self.default_font_family, self.default_font_size))
            return

        self.save_state()
        
        start = self.editor.index(tk.SEL_FIRST)
        end = self.editor.index(tk.SEL_LAST)
        
        curr = start
        while self.editor.compare(curr, "<", end):
            tags = self.editor.tag_names(curr)
            
            # find active style tag
            style_tag = None
            for t in tags:
                if t.startswith("style_"):
                    style_tag = t
                    break
            
            # get current style or use defaults
            if style_tag:
                parts = style_tag.split("_")
                curr_family = parts[1]
                curr_size = int(parts[2])
                curr_weight = parts[3]
                curr_slant = parts[4]
            else:
                curr_family = self.default_font_family
                curr_size = self.default_font_size
                curr_weight = "normal"
                curr_slant = "roman"
                
            # set new style, keeping other styles
            new_family = font_family if font_family is not None else curr_family
            new_size = font_size if font_size is not None else curr_size
            
            if bold is True:
                new_weight = "bold"
            elif bold is False:
                new_weight = "normal"
            else:
                new_weight = curr_weight
                
            if italic is True:
                new_slant = "italic"
            elif italic is False:
                new_slant = "roman"
            else:
                new_slant = curr_slant
                
            new_tag = self.get_or_create_style_tag(new_family, new_size, new_weight, new_slant)
            next_char = self.editor.index(f"{curr} + 1 char")
            
            # update style tags
            if style_tag:
                self.editor.tag_remove(style_tag, curr, next_char)
            self.editor.tag_add(new_tag, curr, next_char)
            
            curr = next_char
            
        self.update_live_stats()

    # undo and redo system
    def get_snapshot(self):
        """get current text, cursor, tags, and images"""
        cursor = self.editor.index(tk.INSERT)
        content = self.editor.get("1.0", "end-1c")
        
        tags = {}
        for tag in self.editor.tag_names():
            if tag == "sel":
                continue
            ranges = self.editor.tag_ranges(tag)
            tag_ranges = []
            for i in range(0, len(ranges), 2):
                tag_ranges.append((self.editor.index(ranges[i]), self.editor.index(ranges[i+1])))
            if tag_ranges:
                tags[tag] = tag_ranges
        
        # capture embedded images
        images = []
        if PIL_AVAILABLE:
            for child in self.editor.winfo_children():
                if isinstance(child, ImageWidget):
                    try:
                        idx = self.editor.index(child)
                        images.append({
                            "index": idx,
                            "width": child.width,
                            "height": child.height,
                            "image": child.original_image
                        })
                    except Exception:
                        pass
                        
        return {"content": content, "cursor": cursor, "tags": tags, "images": images}

    def restore_snapshot(self, snapshot):
        """restore text, cursor, tags, and images"""
        self.editor.delete("1.0", tk.END)
        self.editor.insert("1.0", snapshot["content"])
        
        # clear active tags
        for tag in self.editor.tag_names():
            if tag != "sel":
                self.editor.tag_remove(tag, "1.0", tk.END)
                
        # recreate style tags
        for tag, ranges in snapshot["tags"].items():
            if tag.startswith("style_"):
                parts = tag.split("_")
                if len(parts) == 5:
                    family = parts[1]
                    size = int(parts[2])
                    weight = parts[3]
                    slant = parts[4]
                    self.get_or_create_style_tag(family, size, weight, slant)
            for start, end in ranges:
                self.editor.tag_add(tag, start, end)
                
        # restore embedded images
        if PIL_AVAILABLE:
            for child in self.editor.winfo_children():
                if isinstance(child, ImageWidget):
                    child.destroy()
            for img_data in snapshot.get("images", []):
                try:
                    # check if base64 formatted (loaded from disk)
                    if "base64" in img_data:
                        img_bytes = base64.b64decode(img_data["base64"])
                        img = Image.open(io.BytesIO(img_bytes))
                    else:
                        # raw PIL image object (from in-memory history states)
                        img = img_data["image"]
                        
                    widget = ImageWidget(self.editor, self.editor, img, self, img_data["width"], img_data["height"])
                    self.editor.window_create(img_data["index"], window=widget)
                except Exception as e:
                    print(f"Error restoring image widget: {e}")
                
        self.editor.mark_set(tk.INSERT, snapshot["cursor"])
        self.editor.see(tk.INSERT)
        self.update_live_stats()

    def save_state(self):
        """save state to undo stack"""
        snapshot = self.get_snapshot()
        # avoid saving duplicate states
        if not self.undo_stack or self.undo_stack[-1]["content"] != snapshot["content"] or self.undo_stack[-1]["tags"] != snapshot["tags"] or len(self.undo_stack[-1].get("images", [])) != len(snapshot.get("images", [])):
            self.undo_stack.append(snapshot)
            self.redo_stack.clear()

    def on_key_press(self, event=None):
        """save state before typing changes"""
        if event:
            # ignore modifiers
            if event.keysym in ["Control_L", "Control_R", "Alt_L", "Alt_R", "Shift_L", "Shift_R", "Caps_Lock"]:
                return
            # ignore undo shortcuts
            if event.state & 0x4:
                if event.keysym.lower() in ['z', 'y']:
                    return

        current_time = time.time()
        is_break_or_word_boundary = event and event.keysym in ["space", "Return", "BackSpace", "Delete", "Tab"]
        time_elapsed = current_time - self.last_keystroke_time > 1.5
        
        if is_break_or_word_boundary or time_elapsed or not self.undo_stack:
            self.save_state()
            
        self.last_keystroke_time = current_time

    def undo_action(self, event=None):
        if self.undo_stack:
            current_state = self.get_snapshot()
            self.redo_stack.append(current_state)
            
            prev_state = self.undo_stack.pop()
            self.restore_snapshot(prev_state)
            self.status_bar.config(text="Undo performed.")
        else:
            self.status_bar.config(text="Nothing to undo.")
        return "break"  # prevent native undo

    def redo_action(self, event=None):
        if self.redo_stack:
            current_state = self.get_snapshot()
            self.undo_stack.append(current_state)
            
            next_state = self.redo_stack.pop()
            self.restore_snapshot(next_state)
            self.status_bar.config(text="Redo performed.")
        else:
            self.status_bar.config(text="Nothing to redo.")
        return "break"

    # file actions
    def new_file(self):
        self.editor.delete(1.0, tk.END)
        self.current_file = None
        self.undo_stack.clear()
        self.redo_stack.clear()
        self.root.title("New Document - PyDoc Architect")
        self.status_bar.config(text="New Document created.")
        self.update_live_stats()

    def open_file(self):
        file_path = filedialog.askopenfilename(defaultextension=".docx", 
                                              filetypes=[("Word Documents", "*.docx"), ("PyDoc Files", "*.pydoc"), ("Text Files", "*.txt"), ("All Files", "*.*")])
        if file_path:
            try:
                if file_path.endswith(".docx"):
                    if not DOCX_AVAILABLE:
                        self.show_docx_install_instructions()
                        return
                    self.load_from_docx(file_path)
                elif file_path.endswith(".pydoc"):
                    with open(file_path, "r", encoding="utf-8") as f:
                        snapshot = json.load(f)
                    self.restore_snapshot(snapshot)
                else:
                    with open(file_path, "r", encoding="utf-8") as f:
                        content = f.read()
                    self.editor.delete(1.0, tk.END)
                    self.editor.insert(tk.END, content)
                    self.update_live_stats()
                
                self.current_file = file_path
                self.undo_stack.clear()
                self.redo_stack.clear()
                self.root.title(f"{os.path.basename(file_path)} - PyDoc Architect")
                self.status_bar.config(text=f"Opened: {file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Could not read file: {e}")

    def save_file(self):
        if not self.current_file:
            self.save_as_file()
            return
        self._write_file_data(self.current_file)

    def save_as_file(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                               filetypes=[("Word Documents", "*.docx"), ("PyDoc Files", "*.pydoc"), ("Text Files", "*.txt"), ("All Files", "*.*")])
        if not file_path:
            return
        self.current_file = file_path
        self._write_file_data(self.current_file)

    def _write_file_data(self, file_path):
        """helper to write editor content to disk"""
        try:
            if file_path.endswith(".docx"):
                if not DOCX_AVAILABLE:
                    self.show_docx_install_instructions()
                    return
                self.save_as_docx(file_path)
            elif file_path.endswith(".pydoc"):
                snapshot = self.get_snapshot()
                
                # Base64 encode images so they save in the custom .pydoc JSON format
                serializable_images = []
                for img_data in snapshot.get("images", []):
                    try:
                        buffered = io.BytesIO()
                        fmt = img_data["image"].format if img_data["image"].format else "PNG"
                        img_data["image"].save(buffered, format=fmt)
                        img_str = base64.b64encode(buffered.getvalue()).decode("utf-8")
                        serializable_images.append({
                            "index": img_data["index"],
                            "width": img_data["width"],
                            "height": img_data["height"],
                            "format": fmt,
                            "base64": img_str
                        })
                    except Exception as ex:
                        print(f"Failed to serialize embedded image: {ex}")
                
                serializable_snapshot = snapshot.copy()
                serializable_snapshot["images"] = serializable_images
                
                with open(file_path, "w", encoding="utf-8") as f:
                    json.dump(serializable_snapshot, f, indent=4)
            else:
                content = self.editor.get(1.0, tk.END)
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(content)
            self.root.title(f"{os.path.basename(file_path)} - PyDoc Architect")
            self.status_bar.config(text=f"Saved: {file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Could not save file: {e}")

    # docx import and export with full formatting & sequential image insertion
    def load_from_docx(self, file_path):
        """load docx file with formatting and restore images"""
        doc = docx.Document(file_path)
        self.editor.delete("1.0", tk.END)
        
        for p_idx, p in enumerate(doc.paragraphs):
            if p_idx > 0:
                self.editor.insert(tk.END, "\n")
                
            for run in p.runs:
                # Check for run-contained XML media blips (drawings/embeds)
                embed_id = self._find_embed_id(run.element)
                if embed_id and PIL_AVAILABLE:
                    try:
                        image_part = doc.part.related_parts[embed_id]
                        image_bytes = image_part.image.blob
                        img = Image.open(io.BytesIO(image_bytes))
                        
                        # set layout boundaries cleanly
                        width, height = img.width, img.height
                        if width > 300 or height > 300:
                            aspect = width / height
                            if width > height:
                                width = 300
                                height = int(300 / aspect)
                            else:
                                height = 300
                                width = int(300 * aspect)
                        
                        widget = ImageWidget(self.editor, self.editor, img, self, width, height)
                        self.editor.window_create(tk.END, window=widget)
                        continue  # Skip raw text render if run was exclusively an image
                    except Exception as ex:
                        print(f"Error parsing native DOCX image object: {ex}")
                
                # Render raw text runs normally
                start_index = self.editor.index("end-1c")
                self.editor.insert(tk.END, run.text)
                end_index = self.editor.index("end-1c")
                
                if run.text:
                    # get format styles
                    weight = "bold" if run.bold else "normal"
                    slant = "italic" if run.italic else "roman"
                    family = run.font.name if run.font.name else self.default_font_family
                    size = int(run.font.size.pt) if run.font.size else self.default_font_size
                    
                    tag_name = self.get_or_create_style_tag(family, size, weight, slant)
                    self.editor.tag_add(tag_name, start_index, end_index)
                    
        self.update_live_stats()

    def _find_embed_id(self, element):
        """Recursively search Docx run XML trees for media references"""
        for child in element:
            if child.tag.endswith('blip'):
                for key, val in child.attrib.items():
                    if key.endswith('embed'):
                        return val
            res = self._find_embed_id(child)
            if res:
                return res
        return None

    def save_as_docx(self, file_path):
        """save to docx file with formatting and embedded images sequentially"""
        doc = docx.Document()
        
        # Pull sequential text blocks, images, and markup in order via text widget dumping
        dump_data = self.editor.dump("1.0", "end-1c")
        p = doc.add_paragraph()
        current_tags = set()
        
        for key, value, index in dump_data:
            if key == "tagon":
                if value.startswith("style_"):
                    current_tags.add(value)
            elif key == "tagoff":
                if value in current_tags:
                    current_tags.remove(value)
            elif key == "text":
                parts = value.split("\n")
                for idx, part in enumerate(parts):
                    if idx > 0:
                        p = doc.add_paragraph()
                    if part:
                        run = p.add_run(part)
                        
                        # Search active styling tags
                        style_tag = None
                        for t in current_tags:
                            if t.startswith("style_"):
                                style_tag = t
                                break
                                
                        if style_tag:
                            parts_style = style_tag.split("_")
                            if len(parts_style) == 5:
                                family = parts_style[1]
                                size = int(parts_style[2])
                                weight = parts_style[3]
                                slant = parts_style[4]
                                
                                run.bold = (weight == "bold")
                                run.italic = (slant == "italic")
                                run.font.name = family
                                run.font.size = docx.shared.Pt(size)
                        else:
                            run.font.name = self.default_font_family
                            run.font.size = docx.shared.Pt(self.default_font_size)
                            
            elif key == "window":
                try:
                    widget = self.editor.nametowidget(value)
                    if isinstance(widget, ImageWidget) and PIL_AVAILABLE:
                        img_stream = io.BytesIO()
                        fmt = widget.original_image.format if widget.original_image.format else "PNG"
                        widget.original_image.save(img_stream, format=fmt)
                        img_stream.seek(0)
                        
                        run = p.add_run()
                        # standard screen resolution conversion (96 DPI bounds)
                        width_in = docx.shared.Inches(widget.width / 96.0)
                        height_in = docx.shared.Inches(widget.height / 96.0)
                        run.add_picture(img_stream, width=width_in, height=height_in)
                except Exception as e:
                    print(f"Error packing ImageWidget frame to document stream: {e}")
                    
        doc.save(file_path)

    # font styling helpers
    def change_font(self, event=None):
        family = self.font_family.get()
        try:
            size = int(self.font_size.get())
        except ValueError:
            size = self.default_font_size
        self.apply_formatting(font_family=family, font_size=size)

    def toggle_bold(self):
        try:
            if self.editor.tag_ranges(tk.SEL):
                start = self.editor.index(tk.SEL_FIRST)
                tags = self.editor.tag_names(start)
                is_bold = False
                for t in tags:
                    if t.startswith("style_") and "bold" in t:
                        is_bold = True
                        break
                self.apply_formatting(bold=not is_bold)
        except Exception:
            pass

    def toggle_italic(self):
        try:
            if self.editor.tag_ranges(tk.SEL):
                start = self.editor.index(tk.SEL_FIRST)
                tags = self.editor.tag_names(start)
                is_italic = False
                for t in tags:
                    if t.startswith("style_") and "italic" in t:
                        is_italic = True
                        break
                self.apply_formatting(italic=not is_italic)
        except Exception:
            pass

    def copy_format(self):
        """copy active format style at insert cursor"""
        try:
            cursor = self.editor.index(tk.INSERT)
            tags = self.editor.tag_names(cursor)
            self.format_clipboard = [t for t in tags if t.startswith("style_")]
            if self.format_clipboard:
                self.status_bar.config(text="Format copied.")
            else:
                self.status_bar.config(text="No format found to copy.")
        except Exception as e:
            self.status_bar.config(text=f"Error copying format: {e}")

    def paste_format(self):
        """apply copied format style tags to active selection"""
        if not self.format_clipboard:
            self.status_bar.config(text="No format copied to clipboard.")
            return
        try:
            if self.editor.tag_ranges(tk.SEL):
                self.save_state()
                start = self.editor.index(tk.SEL_FIRST)
                end = self.editor.index(tk.SEL_LAST)
                
                # Clear existing styles
                for t in self.editor.tag_names():
                    if t.startswith("style_"):
                        self.editor.tag_remove(t, start, end)
                # Apply copied styles
                for style_tag in self.format_clipboard:
                    self.editor.tag_add(style_tag, start, end)
                self.status_bar.config(text="Format pasted successfully.")
                self.update_live_stats()
        except Exception as e:
            self.status_bar.config(text=f"Error pasting format: {e}")

    # settings adjustment
    def toggle_dark_mode(self):
        self.is_dark_mode = not self.is_dark_mode
        bg_color = "#1e1e1e" if self.is_dark_mode else "white"
        fg_color = "#ffffff" if self.is_dark_mode else "black"
        insert_color = "#ffffff" if self.is_dark_mode else "black"
        
        self.editor.config(bg=bg_color, fg=fg_color, insertbackground=insert_color)
        
        # update all embedded widgets' backgrounds
        for child in self.editor.winfo_children():
            if isinstance(child, ImageWidget):
                child.config(bg=bg_color)
                child.label.config(bg=bg_color)
                
        btn_text = "Light Mode" if self.is_dark_mode else "Dark Mode"
        self.btn_dark_mode.config(text=btn_text)
        self.status_bar.config(text=f"Switched to {'Dark' if self.is_dark_mode else 'Light'} Mode.")

    def show_docx_install_instructions(self):
        """instructions for installing python-docx"""
        messagebox.showinfo(
            "Word DOCX Integration", 
            "To open and save native Word (.docx) files, please open your system terminal / command prompt and run:\n\n"
            "pip install python-docx\n\n"
            "Then, restart this application! In the meantime, you can save your work using the '.pydoc' or '.txt' formats safely."
        )

    def show_pil_install_instructions(self):
        """instructions for installing pillow"""
        messagebox.showinfo(
            "Image Support", 
            "To paste, move, and resize images, please open your system terminal / command prompt and run:\n\n"
            "pip install Pillow\n\n"
            "Then, restart this application!"
        )

    # image processing actions
    def insert_image_widget(self, img):
        """insert an embedded image widget"""
        self.save_state()
        widget = ImageWidget(self.editor, self.editor, img, self)
        self.editor.window_create(tk.INSERT, window=widget)
        self.update_live_stats()

    def paste_image(self):
        """paste image from system clipboard"""
        if not PIL_AVAILABLE:
            self.show_pil_install_instructions()
            return
        try:
            img = ImageGrab.grabclipboard()
            # grab clipboard filepath list fallback
            if isinstance(img, list) and len(img) > 0 and isinstance(img[0], str) and img[0].lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                img = Image.open(img[0])
            if isinstance(img, Image.Image):
                self.insert_image_widget(img)
                self.status_bar.config(text="Image pasted successfully.")
            else:
                self.status_bar.config(text="No image found in clipboard.")
        except Exception as e:
            self.status_bar.config(text=f"Failed to paste image: {e}")

    def on_paste_event(self, event=None):
        """paste event interceptor"""
        if PIL_AVAILABLE:
            try:
                img = ImageGrab.grabclipboard()
                if isinstance(img, list) and len(img) > 0 and isinstance(img[0], str) and img[0].lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                    img = Image.open(img[0])
                if isinstance(img, Image.Image):
                    self.insert_image_widget(img)
                    self.status_bar.config(text="Image pasted successfully.")
                    return "break"
            except Exception:
                pass

    # bullet lists
    def insert_bullet(self):
        icon = self.bullet_type.get()
        try:
            self.save_state()
            if self.editor.tag_ranges(tk.SEL):
                start_line = int(self.editor.index(tk.SEL_FIRST).split('.')[0])
                end_line = int(self.editor.index(tk.SEL_LAST).split('.')[0])
                for line in range(start_line, end_line + 1):
                    self.editor.insert(f"{line}.0", f"{icon} ")
            else:
                curr_line = self.editor.index(tk.INSERT).split('.')[0]
                self.editor.insert(f"{curr_line}.0", f"{icon} ")
            self.update_live_stats()
        except Exception as e:
            self.status_bar.config(text=f"Failed to insert list item: {e}")

if __name__ == "__main__":
    app_root = tk.Tk()
    app = PyDocApp(app_root)
    app_root.mainloop()